from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name('9life-mag-huong-dan-van-hanh.docx')
GOLD = 'B78A17'
INK = '1B1A18'
MUTED = '666666'
PALE_GOLD = 'F6E8C4'
PALE_GRAY = 'F3F3F3'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ('Heading 1', 16, '2E74B5', 18, 10),
        ('Heading 2', 13, '2E74B5', 14, 7),
        ('Heading 3', 12, '1F4D78', 10, 5),
    ]:
        style = styles[name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run('9LIFE MAG | SỔ TAY VẬN HÀNH')
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Tài liệu nội bộ - cập nhật theo phiên bản hệ thống')
    set_run_font(run, size=8, color=MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('9LIFE MAG')
    set_run_font(run, size=12, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run('Sổ tay giới thiệu và hướng dẫn vận hành')
    set_run_font(run, size=25, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run('Website chính, khu vực tài khoản và CMS quản trị')
    set_run_font(run, size=13, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [2700, 6660])
    labels = [('Mục đích', 'Giới thiệu nhanh hệ thống và quy trình thao tác cơ bản cho đội ngũ vận hành.'), ('Phạm vi', 'Website chính, tài khoản người dùng, cổng nghệ sĩ/đối tác và CMS.'), ('Phiên bản', f'Tài liệu nội bộ - {date.today().strftime("%d/%m/%Y")}')]
    for row, (label, value) in zip(table.rows, labels):
        set_cell_shading(row.cells[0], PALE_GOLD)
        p1 = row.cells[0].paragraphs[0]
        set_run_font(p1.add_run(label), size=10, color=INK, bold=True)
        p2 = row.cells[1].paragraphs[0]
        set_run_font(p2.add_run(value), size=10, color=INK)

    doc.add_page_break()


def heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f'Heading {level}')


def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(item), size=10.5)


def steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(item), size=10.5)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GRAY)
    p = cell.paragraphs[0]
    set_run_font(p.add_run(title + ': '), size=10.5, color=INK, bold=True)
    set_run_font(p.add_run(text), size=10.5, color=INK)


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2400, 3600, 3360])
    headers = ['Khu vực', 'Mục đích', 'Thao tác chính']
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, 'E8EEF5')
        set_run_font(cell.paragraphs[0].add_run(text), size=9.5, color=INK, bold=True)
    for feature, purpose, action in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, [feature, purpose, action]):
            set_run_font(cell.paragraphs[0].add_run(text), size=9.5, color=INK)
    return table


def build_document():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    heading(doc, '1. Tổng quan nền tảng')
    para(doc, '9Life Mag là nền tảng nội dung và cộng đồng nightlife, kết nối tin tức, âm nhạc, nghệ sĩ, outlet/night club, booking và hoạt động tương tác bằng Sao nội bộ.')
    add_feature_table(doc, [
        ('Website chính', 'Không gian công khai để khám phá tin tức, nghệ sĩ, âm nhạc và outlet.', 'Đọc tin, nghe nhạc, theo dõi, vote, đặt bàn, gửi booking và tìm kiếm.'),
        ('Tài khoản', 'Khu vực cá nhân của thành viên.', 'Quản lý Sao, playlist, hoạt động, yêu thích và lịch sử sử dụng.'),
        ('Cổng nghệ sĩ/đối tác', 'Khu vực làm việc riêng theo vai trò.', 'Hoàn thiện hồ sơ, phát hành nhạc, theo dõi booking, quản lý roster hoặc outlet.'),
        ('CMS', 'Trung tâm vận hành dành cho quản trị viên.', 'Duyệt nội dung, quản lý hồ sơ, music catalog, booking, người dùng, Sao và cấu hình tích hợp.'),
    ])
    add_callout(doc, 'Nguyên tắc', 'Dữ liệu nhạy cảm, phê duyệt, phân quyền và cấu hình tích hợp chỉ thực hiện trong CMS bởi người có quyền phù hợp.')

    heading(doc, '2. Hướng dẫn website chính')
    heading(doc, 'Trang chủ', 2)
    para(doc, 'Trang chủ là điểm vào chính, tổng hợp bài viết nổi bật, bảng xếp hạng, nghệ sĩ tiêu biểu và các lựa chọn âm nhạc. Các khối nội dung được thiết kế để mở đúng trang chi tiết hoặc catalog liên quan.')
    bullets(doc, ['Bấm vào ảnh đại diện nghệ sĩ để mở hồ sơ chi tiết.', 'Bấm nút phát để đưa track vào media player cố định ở cuối trang.', 'Vote, tải nhạc và một số tính năng tương tác sẽ yêu cầu đăng nhập và sử dụng Sao.', 'Các danh sách nổi bật có cơ chế phân bổ công bằng để hạn chế lặp lại cùng một hồ sơ hoặc nội dung.'])

    heading(doc, 'Tin tức', 2)
    para(doc, 'Mục Tin tức hiển thị feed bài viết, chuyên đề và các nội dung nổi bật do biên tập viên sắp xếp từ CMS. Người xem có thể mở bài viết chi tiết và tiếp tục đọc các gợi ý liên quan mà không cần quay lại trang danh sách.')
    bullets(doc, ['Dùng thanh tìm kiếm để tra cứu theo từ khóa.', 'Các topic đầu trang được quản trị trong CMS nhằm ưu tiên nội dung theo chủ đề.', 'Trang tự tải thêm bài viết khi cuộn xuống dưới để hỗ trợ theo dõi liên tục.'])

    heading(doc, 'Music', 2)
    para(doc, '9Life Music là catalog nghe nhạc phía người dùng. Media player xuất hiện khi người dùng bấm phát và cho phép quản lý hàng đợi, yêu thích, thêm playlist, chia sẻ, tải xuống hoặc báo cáo sự cố.')
    bullets(doc, ['Bấm Play tại track, nonstop, remix, album hoặc playlist để phát đúng nội dung đã được map.', 'Dùng nút + Playlist để thêm nhạc vào playlist cá nhân.', 'Báo cáo track dùng cho trường hợp nhạc lỗi hoặc khiếu nại bản quyền.', 'Thư viện của bạn là trang riêng để quản lý playlist, lịch sử nghe, lịch sử tải và liên kết chia sẻ.'])

    heading(doc, 'Nghệ sĩ và outlet', 2)
    para(doc, 'Danh mục Nghệ sĩ hỗ trợ lọc theo loại hình như DJ, producer, MC hype, rapper, dancer, photographer và giới tính. Danh mục Đặt bàn hiển thị outlet theo khu vực, địa phương và trạng thái khám phá ngẫu nhiên.')
    bullets(doc, ['Hồ sơ nghệ sĩ có giới thiệu, media kit, link nhạc/video, booking, vote, follow và chia sẻ.', 'Hồ sơ outlet có thông tin không gian, media, sự kiện, vote, đặt bàn và chia sẻ.', 'Khi chưa đăng nhập, thao tác vote/follow sẽ mở yêu cầu đăng nhập thay vì làm mất bố cục trang.'])

    heading(doc, 'Tìm kiếm và liên hệ', 2)
    para(doc, 'Nút tìm kiếm ở menu chính dẫn tới trang kết quả có các tab lọc: Tất cả, Tin tức, Nghệ sĩ, Outlet và Nhạc. Trang Liên hệ tiếp nhận quảng cáo, bản quyền, hợp tác quản lý/agent và các vấn đề khác; yêu cầu sau đó được đưa về CMS để xử lý.')

    heading(doc, '3. Tài khoản người dùng')
    heading(doc, 'Đăng nhập và tạo tài khoản', 2)
    para(doc, 'Trang Tài khoản chia rõ hai tab Đăng nhập và Tạo tài khoản. Hệ thống hỗ trợ email/số điện thoại và có khu vực cấu hình cho các nhà cung cấp xã hội như Google, Facebook.')
    steps(doc, ['Chọn tab Tạo tài khoản và điền thông tin cần thiết.', 'Xác thực theo cấu hình hệ thống, sau đó đăng nhập vào dashboard cá nhân.', 'Nếu quên mật khẩu, dùng luồng Quên mật khẩu để gửi yêu cầu đặt lại.'])

    heading(doc, 'Dashboard User', 2)
    para(doc, 'Dashboard User theo dõi số Sao, nhiệm vụ hằng ngày, lịch sử hoạt động, nghệ sĩ theo dõi, vote, booking và playlist economy.')
    bullets(doc, ['Đăng ký thành công nhận Sao khởi tạo theo cấu hình hiện hành.', 'Mỗi ngày có thể nhận Sao nhiệm vụ; các mốc bonus được hiển thị khi đủ điều kiện.', 'Sao có thể dùng cho vote, nghe nhạc, tải nhạc hoặc quyền lợi premium theo chính sách hệ thống.', 'Playlist có thể tạo từ dashboard hoặc trực tiếp từ nút + Playlist khi nghe nhạc.'])
    add_callout(doc, 'Lưu ý về Sao', 'Sao là đơn vị nội bộ của hệ thống, phục vụ trải nghiệm và kiểm thử hoạt động. Sao không phải tiền tệ, không có giá trị quy đổi thành tiền mặt và không phải công cụ đầu tư.')

    heading(doc, '4. Cổng nghệ sĩ và đối tác')
    heading(doc, 'Dashboard Artist', 2)
    para(doc, 'Dashboard Artist hỗ trợ nghệ sĩ hoàn thiện thông tin trước khi profile được admin duyệt public. Các khu vực chính gồm Hồ sơ, Music, Media, Booking, thông báo và thống kê.')
    bullets(doc, ['Hồ sơ: điền giới thiệu, thể loại, giá booking, ảnh portrait/cover và các liên kết cần thiết.', 'Music: tạo track, nonstop hoặc album; có thể thêm ảnh cover và chọn track có sẵn cho album.', 'Media: tải ảnh và quản lý video/embed từ YouTube, Facebook, Instagram, SoundCloud hoặc Mixcloud theo định dạng hỗ trợ.', 'Booking: xem các yêu cầu thuộc nghệ sĩ theo trạng thái đang tiếp nhận, đã tiếp nhận và hủy.', 'Thông báo: nhận cập nhật về booking, phê duyệt, agent và các yêu cầu liên quan.'])

    heading(doc, 'Đổi Agent', 2)
    para(doc, 'Khối Agent chỉ hiển thị tên Agent hiện tại hoặc Nghệ sĩ tự do. Nghệ sĩ bấm Thay đổi Agent để mở biểu mẫu chọn điểm đến và lý do tùy chọn.')
    bullets(doc, ['Đang có agent và chuyển sang agent mới: agent cũ và agent mới cùng chấp thuận.', 'Đang là Nghệ sĩ tự do và chuyển sang agent mới: chỉ agent mới chấp thuận.', 'Đang có agent và chuyển sang Nghệ sĩ tự do: agent cũ cần chấp thuận; không có agent mới cần duyệt.', 'Nếu agent cũ từ chối, nghệ sĩ có thể gửi khiếu nại để admin xem xét.'])

    heading(doc, 'Dashboard Manager', 2)
    para(doc, 'Manager quản lý danh sách nghệ sĩ đã được admin map vào Agent tương ứng. Dashboard này giúp theo dõi roster, các yêu cầu chuyển agent, profile/release liên quan và thông báo tác nghiệp.')

    heading(doc, 'Dashboard Booking Coordinator', 2)
    para(doc, 'Booking Coordinator chỉ làm việc với các outlet được admin map quyền. Khu vực này tập trung xem và cập nhật trạng thái các booking outlet theo đúng phạm vi được cấp.')
    add_callout(doc, 'Phân quyền', 'Manager và Booking Coordinator không sử dụng CMS. Họ làm việc tại dashboard site chính với dữ liệu bị giới hạn theo mapping do admin phê duyệt.')

    heading(doc, '5. CMS quản trị')
    para(doc, 'CMS là khu vực vận hành nội bộ. Tài khoản CMS cần được admin duyệt quyền; menu và API được giới hạn theo phạm vi quyền thay vì chỉ dựa trên giao diện.')
    add_feature_table(doc, [
        ('Tổng quan', 'Theo dõi sức khỏe vận hành và việc cần xử lý.', 'Xem chỉ số nội dung, booking, user, nghệ sĩ và số liệu Sao theo khoảng thời gian.'),
        ('Bài viết', 'Quản lý tin tức, chuyên mục và chuyên đề.', 'Soạn thảo WYSIWYG, HTML, ảnh, gallery, video, embed, CTA, SEO và publish.'),
        ('Nghệ sĩ', 'Quản lý catalog nghệ sĩ.', 'Tạo/sửa profile, duyệt hồ sơ nháp, xem lĩnh vực, giới tính, agent và trạng thái public.'),
        ('Music', 'Quản lý toàn bộ catalog âm nhạc.', 'Tải file, chọn nghệ sĩ, loại nhạc, album, quyền truy cập, category và map hiển thị.'),
        ('Outlet', 'Quản lý profile night club/outlet.', 'Tạo profile, hình ảnh, video, thông tin venue và cấu hình nhận booking.'),
        ('Booking', 'Theo dõi yêu cầu từ site chính.', 'Lọc Tất cả/Nghệ sĩ/Đặt bàn/Liên hệ; tiếp nhận, hủy, cấu hình thông báo Telegram.'),
        ('Người dùng', 'Quản lý tài khoản và vai trò đối tác.', 'Duyệt quyền, map Manager với Agent, map Booking Coordinator với Outlet.'),
        ('Sao & thanh toán', 'Quản lý gói Sao và đối soát.', 'Cấu hình gateway, duyệt/từ chối nạp Sao, kiểm tra ví thành viên.'),
        ('API & bảo mật', 'Quản lý tích hợp kỹ thuật.', 'Cấu hình R2, social login, email, payment và các khóa bí mật qua biến môi trường.'),
    ])

    heading(doc, '6. Quy trình vận hành cơ bản')
    heading(doc, 'Duyệt profile nghệ sĩ mới', 2)
    steps(doc, ['Mở CMS > Nghệ sĩ và lọc các profile cần rà soát.', 'Kiểm tra thông tin nhận diện, ảnh, mô tả, liên kết nhạc/video và giá booking nếu có.', 'Chỉnh sửa khi cần, sau đó chuyển trạng thái sang public để profile hiển thị trên site chính.', 'Nếu cần từ chối hoặc yêu cầu bổ sung, dùng thông báo/ghi chú để phản hồi rõ ràng cho nghệ sĩ.'])

    heading(doc, 'Đăng bài viết', 2)
    steps(doc, ['Mở CMS > Bài viết và tạo bài mới.', 'Chọn chuyên mục/chuyên đề, nhập tiêu đề, mô tả, ảnh đại diện và nội dung.', 'Dùng toolbar để chèn ảnh, gallery, video, embed, CTA hoặc chuyển sang HTML khi cần.', 'Rà soát preview, SEO và liên kết trước khi publish hoặc hẹn lịch xuất bản.'])

    heading(doc, 'Tiếp nhận booking', 2)
    steps(doc, ['Mở CMS > Booking và chọn đúng nhóm: Booking nghệ sĩ, Đặt bàn hoặc Liên hệ.', 'Mở chi tiết để kiểm tra toàn bộ thông tin người dùng đã gửi.', 'Chọn Tiếp nhận hoặc Hủy; hệ thống chuyển yêu cầu sang tab trạng thái tương ứng.', 'Cấu hình Telegram tại kênh tổng và profile liên quan để nhận thông báo, soundcheck/check-in hoặc nhắc việc theo server time.'])

    heading(doc, 'Quản lý nhạc', 2)
    steps(doc, ['Mở CMS > Music > Upload nhạc.', 'Chọn nghệ sĩ nếu nội dung thuộc một profile; để trống nếu là nội dung hệ thống.', 'Chọn loại Track, Nonstop, Remix, Playlist hoặc Album và category phù hợp.', 'Thiết lập quyền truy cập, map vị trí hiển thị và kiểm tra cover trước khi publish.', 'Với R2 và worker xử lý media, chỉ dùng khóa cấu hình trong CMS/API hoặc biến môi trường, không đưa khóa vào nội dung bài viết.'])

    heading(doc, '7. Lưu ý bảo mật và vận hành')
    bullets(doc, ['Không chia sẻ mật khẩu, token Telegram, khóa R2, OAuth secret hoặc thông tin payment qua tin nhắn công khai.', 'Tài khoản CMS chỉ cấp theo nguyên tắc đủ dùng; định kỳ rà soát và thu hồi quyền không còn cần thiết.', 'Manager/Booking Coordinator phải được admin map đúng Agent/Outlet trước khi thấy dữ liệu nghiệp vụ.', 'Mọi thao tác nạp Sao, phê duyệt booking, đổi agent hoặc publish nội dung cần giữ lịch sử và thông báo liên quan.', 'Khi đưa lên production, dùng HTTPS, biến môi trường riêng, database thật, sao lưu định kỳ và giám sát lỗi/uptime.', 'Không coi việc chặn F12 là biện pháp bảo mật. Bảo mật thực sự nằm ở kiểm soát API, phân quyền server-side, xác thực và quản lý secret.'])

    heading(doc, '8. Danh sách kiểm tra trước khi vận hành thật')
    bullets(doc, ['Đã cấu hình database production và backup tự động.', 'Đã cấu hình Cloudflare R2/CDN, domain public và CORS phù hợp.', 'Đã đặt toàn bộ secret trong biến môi trường production, không commit vào source.', 'Đã cấu hình Google/Facebook login, email gửi reset mật khẩu và payment gateway theo tài khoản doanh nghiệp.', 'Đã cấu hình channel Telegram tổng và channel riêng theo profile/outlet khi cần.', 'Đã phân quyền CMS, Manager, Booking Coordinator và thực hiện test bằng tài khoản từng vai trò.', 'Đã kiểm tra luồng đăng ký, đăng nhập, quên mật khẩu, upload media, booking, vote, Sao và phát nhạc.', 'Đã tạo lịch backup source code và database, đồng thời kiểm tra khả năng khôi phục.'])

    add_callout(doc, 'Hỗ trợ nội bộ', 'Khi phát hiện lỗi, hãy ghi rõ trang đang thao tác, tài khoản/role, thời gian, các bước tái hiện và ảnh chụp màn hình. Thông tin đầy đủ giúp đội kỹ thuật xử lý nhanh và an toàn hơn.')
    doc.save(OUTPUT)


if __name__ == '__main__':
    build_document()
